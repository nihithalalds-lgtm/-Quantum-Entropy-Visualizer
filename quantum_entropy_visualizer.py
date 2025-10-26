"""
Quantum Entropy Visualizer — Streamlit App
File: quantum_entropy_visualizer.py

This single-file Streamlit app implements:
- A simple quantum-density-matrix simulator for a small number of qubits
- Time evolution with controllable noise and unitary rotations
- Calculation of Von Neumann entropy over time
- Interactive GUI controls (sliders, buttons)
- Graph generation (plotly) and image export
- "Export Report" -> generates a Word document (.docx) with placeholders for figures and written content
- Helpful comments and sections where you can leave space for pictures in your final report

Requirements (put into requirements.txt):
streamlit
numpy
scipy
plotly
pandas
python-docx

Run:
1. pip install -r requirements.txt
2. streamlit run quantum_entropy_visualizer.py

Notes:
- This is an educational simulator, not a real quantum mechanics engine. It uses small random/unitary matrices for demonstration.
- The code is commented in sections: Simulator, GUI, Export, Git instructions placeholder.

"""

import streamlit as st
import numpy as np
from scipy.linalg import expm
import pandas as pd
import plotly.express as px
from io import BytesIO
from docx import Document
from docx.shared import Inches
import base64
import time

# ---------------------------
# Helper: Quantum utilities
# ---------------------------

def random_density_matrix(dim):
    A = (np.random.randn(dim, dim) + 1j * np.random.randn(dim, dim)) / np.sqrt(2)
    rho = A @ A.conj().T
    rho = rho / np.trace(rho)
    return rho


def von_neumann_entropy(rho, eps=1e-12):
    eigenvals = np.real_if_close(np.linalg.eigvalsh(rho))
    eigenvals = np.clip(eigenvals, eps, None)
    S = -np.sum(eigenvals * np.log2(eigenvals))
    return float(S)


def single_qubit_rotation(theta, phi=0.0):
    return np.array([[np.cos(theta/2), -np.exp(1j*phi)*np.sin(theta/2)],
                     [np.exp(1j*phi)*np.sin(theta/2), np.cos(theta/2)]])


def tensor_product(unitaries):
    U = unitaries[0]
    for u in unitaries[1:]:
        U = np.kron(U, u)
    return U


def apply_unitary(rho, U):
    return U @ rho @ U.conj().T

# ---------------------------
# Simulator: time evolution
# ---------------------------

def simulate_entropy(num_qubits=2, time_steps=50, rotation_strength=0.2, noise_level=0.01, seed=None):
    if seed is not None:
        np.random.seed(seed)

    dim = 2 ** num_qubits
    rho = random_density_matrix(dim)

    entropies = []
    times = np.arange(time_steps)

    for t in times:
        unitaries = []
        for q in range(num_qubits):
            theta = rotation_strength * (1 + 0.3 * np.sin(0.2 * t + q)) + 0.02 * np.random.randn()
            phi = 0.1 * np.random.randn()
            unitaries.append(single_qubit_rotation(theta, phi))
        U = tensor_product(unitaries)
        rho = apply_unitary(rho, U)
        maximally_mixed = np.eye(dim) / dim
        rho = (1 - noise_level) * rho + noise_level * maximally_mixed
        rho = rho / np.trace(rho)
        S = von_neumann_entropy(rho)
        entropies.append(S)

    df = pd.DataFrame({'time': times, 'entropy_bits': entropies})
    return df

# ---------------------------
# Export helpers
# ---------------------------

def fig_to_image_bytes(fig):
    return fig.to_image(format='png')


def create_word_report(df, fig_png_bytes, project_title="Quantum Entropy Visualizer Report", author="Author Name"):
    doc = Document()
    doc.add_heading(project_title, level=1)
    doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    doc.add_heading('1. Summary', level=2)
    doc.add_paragraph('Write your executive summary here. Leave space for figures and observations.')

    doc.add_heading('2. Simulation Parameters', level=2)
    doc.add_paragraph('Describe parameters used in the run. Example: number of qubits, noise level, rotation strength, time steps.')

    doc.add_heading('3. Results', level=2)
    doc.add_paragraph('Attach the entropy-over-time figure below:')
    image_stream = BytesIO(fig_png_bytes)
    doc.add_picture(image_stream, width=Inches(6))
    doc.add_paragraph('Figure 1: Entropy (bits) vs simulation time. \n\n')

    doc.add_heading('4. Images / Screenshots (Placeholders)', level=2)
    doc.add_paragraph('Insert your images and captions here. (Leave blank in automated export.)')

    doc.add_heading('5. Incident Response Notes', level=2)
    doc.add_paragraph('Use this section to describe the incident scenario you simulated and the response timeline.\n')

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# ---------------------------
# Streamlit GUI
# ---------------------------

st.set_page_config(page_title='Quantum Entropy Visualizer', layout='wide')
st.title('🔮 Quantum Entropy Visualizer — Incident Response Simulator')

with st.sidebar:
    st.header('Simulator Controls')
    num_qubits = st.slider('Number of qubits', 1, 4, 2)
    time_steps = st.slider('Time steps', 10, 500, 100)
    rotation_strength = st.slider('Rotation strength (rad)', 0.0, 1.5, 0.25, step=0.01)
    noise_level = st.slider('Noise / decoherence (mixing fraction)', 0.0, 0.5, 0.02, step=0.005)
    random_seed = st.number_input('Random seed (0 = random)', min_value=0, value=0)
    run_sim = st.button('Run Simulation')
    export_word = st.button('Export Report (Word)')

st.markdown('---')
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader('Entropy over time')
    placeholder_fig = st.empty()

with col2:
    st.subheader('Incident Response Panel')
    incident_title = st.text_input('Incident title', value='Suspicious entropy spike')
    incident_description = st.text_area('Describe the incident and response steps (leave space for images):', value='1) Detect spike\n2) Isolate system\n3) Analyse density matrices\n4) Recover backups')
    st.markdown('**Suggested response checklist:**')
    st.write('- Capture system snapshot\n- Save logs and visualizations\n- Notify stakeholders\n- Run containment playbook')

if run_sim:
    seed = None if random_seed == 0 else int(random_seed)
    df = simulate_entropy(num_qubits=num_qubits, time_steps=time_steps, rotation_strength=rotation_strength, noise_level=noise_level, seed=seed)
    fig = px.line(df, x='time', y='entropy_bits', title='Von Neumann Entropy over Time', labels={'entropy_bits': 'Entropy (bits)', 'time': 'Time step'})
    fig.update_layout(height=450)
    placeholder_fig.plotly_chart(fig, use_container_width=True)

    st.subheader('Simulation data (sample)')
    st.dataframe(df.head(20))
    st.markdown('**Quick stats**')
    st.write(df.describe())

    if export_word:
        st.info('Preparing Word report...')
        png_bytes = fig_to_image_bytes(fig)
        word_bytes = create_word_report(df, png_bytes, project_title=incident_title, author='Your Name')
        b64 = base64.b64encode(word_bytes).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="quantum_entropy_report.docx">Download Word report</a>'
        st.markdown(href, unsafe_allow_html=True)
        st.success('Report ready for download')
else:
    st.info('Adjust parameters in the sidebar and press **Run Simulation** to generate entropy plots.')

st.markdown('---')
st.header('Appendix — How to push this project to GitHub')
st.markdown("""
1. Create a new folder and place `quantum_entropy_visualizer.py` and `requirements.txt` in it.
2. `git init`
3. `git add .`
4. `git commit -m "Initial commit: Quantum Entropy Visualizer"`
5. Create a GitHub repo and connect:
   `git remote add origin git@github.com:yourusername/quantum-entropy-visualizer.git`
   `git branch -M main`
   `git push -u origin main`
6. Add README.md and LICENSE if desired.
""")

st.caption("""End of automated simulator. 
Placeholders for images in the Word report are left intentionally so you can add screenshots and incident artefacts.""")